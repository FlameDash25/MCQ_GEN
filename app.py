import os
from pathlib import Path
from flask import Flask, render_template, request, send_file, flash, redirect
import pdfplumber
import docx
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from werkzeug.utils import secure_filename
import google.generativeai as genai
from fpdf import FPDF
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from datetime import datetime
import logging
import pytesseract
from PIL import Image
import openpyxl
from openpyxl.styles import Font, Color, Alignment, PatternFill
from openpyxl.utils import get_column_letter
import arabic_reshaper
from bidi.algorithm import get_display
import unicodedata

# ============================
# 🔑 API Key Configuration
# ============================
os.environ["GOOGLE_API_KEY"] = "AIzaSyCNnxJlPX7lo8AF_D6sysn-QL3t7awcH1M"
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
model = genai.GenerativeModel("models/gemini-1.5-pro")

# ============================
# ⚙ Flask App Configuration
# ============================
app = Flask(__name__)
app.secret_key = "supersecretkey"  # Required for flash messages
UPLOAD_FOLDER, RESULTS_FOLDER = Path("uploads"), Path("results")
ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}

app.config.update(UPLOAD_FOLDER=UPLOAD_FOLDER, RESULTS_FOLDER=RESULTS_FOLDER)
executor = ThreadPoolExecutor(max_workers=4)

UPLOAD_FOLDER.mkdir(exist_ok=True)
RESULTS_FOLDER.mkdir(exist_ok=True)


# ============================
# 🌐 Language Support Functions
# ============================
def is_rtl(text):
    """Check if text contains RTL characters (Arabic, Hebrew, etc.)"""
    rtl_scripts = [
        "Arabic",
        "Hebrew",
        "Syriac",
        "Thaana",
        "NKo",
        "Samaritan",
        "Mandaic",
        "Devanagari",
        "Bengali",
    ]
    for char in text:
        try:
            script = unicodedata.name(char).split()[0]
            if script in rtl_scripts:
                return True
        except ValueError:
            continue
    return False


def process_rtl_text(text):
    """Process RTL text for proper display (Arabic reshaping and bidi algorithm)"""
    try:
        reshaped_text = arabic_reshaper.reshape(text)
        return get_display(reshaped_text)
    except:
        return text


# ============================
# 📂 Helper Functions
# ============================
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_file(file_path):
    ext = file_path.suffix.lower()
    try:
        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif ext == ".docx":
            return "\n".join(para.text for para in docx.Document(file_path).paragraphs)
        elif ext == ".txt":
            return file_path.read_text(encoding="utf-8")
    except Exception as e:
        logging.error(f"Error extracting text: {str(e)}")
        return None
    return ""


def analyze_image(image_path):
    """Extract text from an image using Tesseract OCR with language detection"""
    try:
        pytesseract.pytesseract.tesseract_cmd = (
            r"C:/Program Files/Tesseract-OCR/tesseract.exe"  # Windows
        )
        img = Image.open(image_path)

        # Try to detect language automatically
        langs = {
            "eng": "English",
            "ara": "Arabic",
            "chi_sim": "Chinese",
            "chi_tra": "Chinese",
            "fra": "French",
            "deu": "German",
            "hin": "Hindi",
            "jpn": "Japanese",
            "kor": "Korean",
            "rus": "Russian",
            "spa": "Spanish",
        }

        # First try with English
        text = pytesseract.image_to_string(img, lang="eng")
        if not text.strip():
            # If no English text found, try with other languages
            for lang_code in langs:
                if lang_code != "eng":
                    try:
                        text = pytesseract.image_to_string(img, lang=lang_code)
                        if text.strip():
                            break
                    except:
                        continue

        return text.strip() if text else ""
    except Exception as e:
        logging.error(f"Error analyzing image: {str(e)}")
        return ""


def extract_text_and_images_from_pdf(file_path, start_page=None, end_page=None):
    """Extract text and images from a PDF file within a specified page range."""
    text = ""
    images = []
    with pdfplumber.open(file_path) as pdf:
        start_page = start_page - 1 if start_page else 0
        end_page = end_page if end_page else len(pdf.pages)

        for page_num, page in enumerate(pdf.pages):
            if start_page <= page_num < end_page:
                text += page.extract_text() or ""
                for image in page.images:
                    image_path = RESULTS_FOLDER / f"image_{len(images)}.png"
                    with open(image_path, "wb") as img_file:
                        img_file.write(image["stream"].get_data())
                    images.append(image_path)
    return text, images


def generate_mcqs_from_text_and_images(
    text, images, num_questions, difficulty, target_language="none"
):
    """Generate MCQs from text and images using Gemini API with optional translation."""
    image_texts = []
    for image in images:
        image_text = analyze_image(image)
        image_texts.append(image_text)

    full_content = text + "\n".join(image_texts)

    language_instructions = {
        "none": "",  # No translation
        "english": " - Translate all questions, options, and answers to English.",
        "hindi": " - Translate all questions, options, and answers to Hindi (Devanagari script).",
        "spanish": " - Translate all questions, options, and answers to Spanish.",
        "french": " - Translate all questions, options, and answers to French.",
        "german": " - Translate all questions, options, and answers to German.",
        "japanese": " - Translate all questions, options, and answers to Japanese (use Kanji where appropriate).",
        "korean": " - Translate all questions, options, and answers to Korean (Hangul script).",
        "chinese": " - Translate all questions, options, and answers to Simplified Chinese.",
        "arabic": " - Translate all questions, options, and answers to Arabic (right-to-left script).",
        "russian": " - Translate all questions, options, and answers to Russian (Cyrillic script).",
        "portuguese": " - Translate all questions, options, and answers to Portuguese.",
        "italian": " - Translate all questions, options, and answers to Italian.",
        "dutch": " - Translate all questions, options, and answers to Dutch.",
        "swedish": " - Translate all questions, options, and answers to Swedish.",
        "turkish": " - Translate all questions, options, and answers to Turkish.",
        "polish": " - Translate all questions, options, and answers to Polish.",
        "ukrainian": " - Translate all questions, options, and answers to Ukrainian.",
        "vietnamese": " - Translate all questions, options, and answers to Vietnamese.",
        "thai": " - Translate all questions, options, and answers to Thai.",
        "greek": " - Translate all questions, options, and answers to Greek.",
        "hebrew": " - Translate all questions, options, and answers to Hebrew.",
        "farsi": " - Translate all questions, options, and answers to Farsi (Persian).",
    }

    difficulty_instructions = {
        "easy": """
        - Generate simple and straightforward questions suitable for beginners.
        - Focus on basic concepts and facts.
        - Avoid complex or ambiguous wording.
        """,
        "medium": """
        - Generate moderately challenging questions suitable for intermediate learners.
        - Include questions that require understanding of relationships between concepts.
        - Use some application-based or scenario-based questions.
        """,
        "hard": """
        - Generate complex and challenging questions suitable for advanced learners.
        - Include questions that require critical thinking, analysis, or synthesis.
        - Use application-based, scenario-based, or problem-solving questions.
        """,
    }

    prompt = f"""
    You are an expert MCQ generator. Your task is to generate {num_questions} multiple-choice questions (MCQs) based on the following text and any diagrams or visual content:

    ### Instructions:
    - Difficulty Level: {difficulty.capitalize()}
    {difficulty_instructions[difficulty]}
    {language_instructions.get(target_language, "")}
    - Ensure questions are **clear, relevant, and unambiguous**.
    - Provide **4 distinct answer choices** for each question.
    - Indicate the **correct answer** at the end of each question.
    - Carefully analyze both the text and any diagrams or visual content to ensure the questions and distractors are accurate and relevant.
    - Do not generate question numbers anywhere.
    - In Question don't include phrases like "According to the text" anywhere or "In the Text"
    - Do not include any HTML tags (like <sup>, <sub>, <b>, etc.) in the output.

    ### Text Content:
    '{full_content}'

    ### Output Format:
    For each MCQ, use the following format:
    ## MCQ
    Question: [Your question here]
    A) [Option A]
    B) [Option B]
    C) [Option C]
    D) [Option D]
    Correct Answer: [Correct option letter]

    Now, generate {num_questions} MCQs based on the provided text and difficulty level.
    """
    try:
        response = model.generate_content(prompt)
        return response.text.strip() if response else "No MCQs generated."
    except Exception as e:
        logging.error(f"Error generating MCQs: {str(e)}")
        return None


def save_text_file(content, filename):
    """Save the generated MCQs to a text file with proper encoding."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_path = RESULTS_FOLDER / filename

    # Process RTL text if needed
    if is_rtl(content):
        content = process_rtl_text(content)

    content_with_timestamp = (
        f"Generated on: {timestamp}\n"
        f"Note: For best readability, use a Unicode-compatible font.\n\n"
        f"{content}"
    )
    try:
        file_path.write_text(content_with_timestamp, encoding="utf-8")
        return file_path
    except Exception as e:
        logging.error(f"Error saving text file: {str(e)}")
        return None


def generate_pdf(mcqs, filename):
    """Generate a PDF file from the MCQs with full language support."""
    try:
        pdf = FPDF()
        pdf.add_page()

        # Add Noto fonts for multilingual support
        font_path = "fonts/NotoSans-Regular.ttf"
        font_bold_path = "fonts/NotoSans-Bold.ttf"

        # Check if we need RTL support
        rtl_needed = is_rtl(mcqs)

        # Add main font
        pdf.add_font("NotoSans", "", font_path, uni=True)
        pdf.add_font("NotoSans", "B", font_bold_path, uni=True)

        # Set the default font
        pdf.set_font("NotoSans", size=12)

        # Add title
        title = "Generated MCQs"
        if rtl_needed:
            title = process_rtl_text(title)
        pdf.cell(0, 10, title, ln=True, align="C")

        # Clean content
        mcqs = mcqs.replace("\r", "").strip()

        for i, mcq in enumerate(mcqs.split("## MCQ"), start=0):
            if mcq.strip():
                lines = mcq.strip().split("\n")
                correct_answer = None

                # Extract the correct answer
                for line in lines:
                    if "Correct Answer:" in line:
                        correct_answer = line.split(":")[-1].strip()
                        break

                # Extract question text
                question_text = (
                    lines[0].split("Question:", 1)[-1].strip()
                    if "Question:" in lines[0]
                    else lines[0].strip()
                )

                # Process RTL text if needed
                if rtl_needed:
                    question_text = process_rtl_text(question_text)

                # Add question
                pdf.set_font("NotoSans", "B", 12)
                question = f"Q{i}: {question_text}"
                pdf.multi_cell(0, 10, question)
                pdf.ln(2)

                # Add options
                pdf.set_font("NotoSans", size=12)
                for line in lines[1:]:
                    if "Correct Answer:" in line:
                        continue

                    if line.strip():
                        option_text = line.strip()
                        if option_text[:2].isdigit() and option_text[2] == ")":
                            option_text = option_text[3:].strip()

                        # Process RTL text if needed
                        if rtl_needed:
                            option_text = process_rtl_text(option_text)

                        # Highlight correct answer
                        if correct_answer and line.strip().startswith(
                            correct_answer + ")"
                        ):
                            pdf.set_text_color(0, 128, 0)
                            pdf.set_font("NotoSans", "B", 12)
                        else:
                            pdf.set_text_color(0, 0, 0)
                            pdf.set_font("NotoSans", size=12)

                        pdf.multi_cell(0, 10, option_text)
                        pdf.ln(2)

                # Add correct answer
                pdf.set_text_color(0, 128, 0)
                pdf.set_font("NotoSans", "B", 12)
                correct = f"Correct Answer: {correct_answer}"
                if rtl_needed:
                    correct = process_rtl_text(correct)
                pdf.multi_cell(0, 10, correct)
                pdf.ln(10)

                # Reset text color
                pdf.set_text_color(0, 0, 0)

        # Add timestamp
        pdf.set_font("NotoSans", size=10)
        timestamp = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        if rtl_needed:
            timestamp = process_rtl_text(timestamp)
        pdf.cell(0, 10, timestamp, ln=True)

        # Save PDF
        pdf_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
        pdf.output(pdf_path)
        return pdf_path
    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        return None


def generate_docx(mcqs, filename):
    """Generate a DOCX file with full language support."""
    try:
        doc = Document()

        # Set default style
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial Unicode MS"  # Broad Unicode support
        font.size = Pt(12)

        # Check for RTL languages
        rtl_needed = is_rtl(mcqs)

        # Enable RTL if needed
        if rtl_needed:
            doc.styles["Normal"]._element.rPr.add(qn("w:bidi"), "on")
            doc.styles["Normal"]._element.rPr.add(qn("w:rtl"), "on")

        # Add title
        title = doc.add_paragraph("Generated MCQs")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(14)

        # Add timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on: {timestamp}\n").alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        # Process MCQs
        for i, mcq in enumerate(mcqs.split("## MCQ"), start=0):
            if mcq.strip():
                lines = mcq.strip().split("\n")
                correct_answer = None

                # Extract correct answer
                for line in lines:
                    if "Correct Answer:" in line:
                        correct_answer = line.split(":")[-1].strip()
                        break

                # Extract question
                question_text = (
                    lines[0].split("Question:", 1)[-1].strip()
                    if "Question:" in lines[0]
                    else lines[0].strip()
                )

                # Process RTL text if needed
                if rtl_needed:
                    question_text = process_rtl_text(question_text)

                # Add question
                question_para = doc.add_paragraph()
                question_run = question_para.add_run(f"Q{i}: {question_text}")
                question_run.bold = True
                if rtl_needed:
                    question_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                # Add options
                for line in lines[1:]:
                    if "Correct Answer:" in line:
                        continue

                    if line.strip():
                        option_text = line.strip()
                        if option_text[:2].isdigit() and option_text[2] == ")":
                            option_text = option_text[3:].strip()

                        # Process RTL text if needed
                        if rtl_needed:
                            option_text = process_rtl_text(option_text)

                        option_para = doc.add_paragraph(option_text)
                        if rtl_needed:
                            option_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                        # Highlight correct answer
                        if correct_answer and line.strip().startswith(
                            correct_answer + ")"
                        ):
                            for run in option_para.runs:
                                run.font.color.rgb = RGBColor(0, 128, 0)
                                run.bold = True

                # Add correct answer
                correct_text = f"Correct Answer: {correct_answer}"
                if rtl_needed:
                    correct_text = process_rtl_text(correct_text)

                correct_para = doc.add_paragraph(correct_text)
                if rtl_needed:
                    correct_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                for run in correct_para.runs:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    run.bold = True

                doc.add_paragraph("\n")  # Add spacing

        # Save DOCX
        docx_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
        doc.save(docx_path)
        return docx_path
    except Exception as e:
        logging.error(f"Error generating DOCX: {str(e)}")
        return None


def generate_excel(mcqs, filename):
    """Generate an Excel file with full language support."""
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "MCQs"

        # Add headers
        headers = [
            "Question",
            "Option A",
            "Option B",
            "Option C",
            "Option D",
            "Correct Answer",
        ]

        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(
                start_color="DDDDDD", end_color="DDDDDD", fill_type="solid"
            )
            cell.alignment = Alignment(horizontal="center")

        # Set column widths
        column_widths = [60, 30, 30, 30, 30, 15]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width

        # Process MCQs
        row_num = 2
        for mcq in mcqs.split("## MCQ"):
            if mcq.strip():
                lines = mcq.strip().split("\n")
                question = ""
                options = {"A": "", "B": "", "C": "", "D": ""}
                correct_answer = ""

                # Extract question and options
                for line in lines:
                    if line.startswith("Question:"):
                        question = line.split("Question:", 1)[-1].strip()
                    elif line.startswith("A)"):
                        options["A"] = line.split("A)", 1)[-1].strip()
                    elif line.startswith("B)"):
                        options["B"] = line.split("B)", 1)[-1].strip()
                    elif line.startswith("C)"):
                        options["C"] = line.split("C)", 1)[-1].strip()
                    elif line.startswith("D)"):
                        options["D"] = line.split("D)", 1)[-1].strip()
                    elif line.startswith("Correct Answer:"):
                        correct_answer = line.split("Correct Answer:", 1)[-1].strip()

                # Write to Excel
                ws.cell(row=row_num, column=1, value=question).alignment = Alignment(
                    wrap_text=True
                )
                ws.cell(row=row_num, column=2, value=options["A"]).alignment = (
                    Alignment(wrap_text=True)
                )
                ws.cell(row=row_num, column=3, value=options["B"]).alignment = (
                    Alignment(wrap_text=True)
                )
                ws.cell(row=row_num, column=4, value=options["C"]).alignment = (
                    Alignment(wrap_text=True)
                )
                ws.cell(row=row_num, column=5, value=options["D"]).alignment = (
                    Alignment(wrap_text=True)
                )

                # Format correct answer
                correct_cell = ws.cell(row=row_num, column=6, value=correct_answer)
                correct_cell.font = Font(color="008000", bold=True)
                correct_cell.alignment = Alignment(horizontal="center")

                row_num += 1

        # Add timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        ws.cell(row=row_num + 1, column=1, value=f"Generated on: {timestamp}").font = (
            Font(italic=True)
        )

        # Save Excel
        excel_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
        wb.save(excel_path)
        return excel_path
    except Exception as e:
        logging.error(f"Error generating Excel file: {str(e)}")
        return None


# ============================
# 📌 Flask Routes
# ============================
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/contact")
def contact():
    return render_template("contact.html")


@app.route("/generate", methods=["POST"])
def generate_mcqs():
    if "file" not in request.files:
        flash("No file uploaded", "error")
        return redirect(request.url)

    file = request.files["file"]
    if file.filename == "":
        flash("No file selected", "error")
        return redirect(request.url)

    if not allowed_file(file.filename):
        flash("Invalid file format. Allowed formats: PDF, TXT, DOCX", "error")
        return redirect(request.url)

    filename = secure_filename(file.filename)
    file_path = UPLOAD_FOLDER / filename
    file.save(file_path)

    # Get page range
    start_page = request.form.get("start_page", type=int)
    end_page = request.form.get("end_page", type=int)

    if start_page and end_page and start_page > end_page:
        flash("Start page cannot be greater than end page", "error")
        return redirect(request.url)

    # Extract text and images
    if file_path.suffix.lower() == ".pdf":
        text, images = extract_text_and_images_from_pdf(file_path, start_page, end_page)
    else:
        text = extract_text_from_file(file_path)
        images = []

    if not text:
        flash("Could not extract text from file", "error")
        return redirect(request.url)

    try:
        num_questions = int(request.form["num_questions"])
        if num_questions <= 0:
            flash("Number of questions must be greater than zero", "error")
            return redirect(request.url)
    except ValueError:
        flash("Invalid number of questions", "error")
        return redirect(request.url)

    # Get difficulty and language
    difficulty = request.form["difficulty"]
    target_language = request.form.get("language", "none")

    # Generate MCQs
    mcqs = generate_mcqs_from_text_and_images(
        text, images, num_questions, difficulty, target_language
    )
    if not mcqs:
        flash("Failed to generate MCQs", "error")
        return redirect(request.url)

    base_filename = filename.rsplit(".", 1)[0]
    txt_filename = f"generated_mcqs_{base_filename}.txt"
    pdf_filename = f"generated_mcqs_{base_filename}.pdf"
    docx_filename = f"generated_mcqs_{base_filename}.docx"
    excel_filename = f"generated_mcqs_{base_filename}.xlsx"

    # Generate all file formats
    if not save_text_file(mcqs, txt_filename):
        flash("Failed to save text file", "error")
        return redirect(request.url)

    if not generate_pdf(mcqs, pdf_filename):
        flash("Failed to generate PDF", "error")
        return redirect(request.url)

    if not generate_docx(mcqs, docx_filename):
        flash("Failed to generate DOCX", "error")
        return redirect(request.url)

    if not generate_excel(mcqs, excel_filename):
        flash("Failed to generate Excel file", "error")
        return redirect(request.url)

    return render_template(
        "results.html",
        mcqs=mcqs,
        txt_filename=txt_filename,
        pdf_filename=pdf_filename,
        docx_filename=docx_filename,
        excel_filename=excel_filename,
    )


@app.route("/download/<filename>")
def download_file(filename):
    file_path = RESULTS_FOLDER / filename
    if not file_path.exists():
        flash("File not found", "error")
        return redirect("/")
    return send_file(file_path, as_attachment=True)


# ============================
# 🚀 Run the Flask App
# ============================
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    app.run(host="0.0.0.0", port=5000, debug=True, threaded=True)
